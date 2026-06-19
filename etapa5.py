"""
etapa5.py — Etapa 5: Comparação Técnica (one-pager).

O que faz:
1. NÃO é análise nova — é SÍNTESE do que as Etapas 2, 3 e 4 já produziram:
   Baseline x Edital x Propostas Técnicas, ótica técnica, sem preço.
2. Pede ao Claude uma consolidação: matriz requisito x fornecedor, gaps
   mandatórios por fornecedor, cruzamento de inclusões/exclusões de escopo,
   e uma leitura do que isso significa pra decisão (ainda sem custo).
3. Com 1 proposta só, a "comparação" degenera pra um resumo de conformidade
   único — sem ranking entre fornecedores. O prompt já assume isso e o
   código registra a limitação na Memória de Premissas.
4. Grava em estudo.comparacao_tecnica.
5. Gera as duas primeiras saídas formatadas do projeto: Word (one-pager) e
   Excel (matriz detalhada). PPT na marca A&M: pendente layout do usuário.
"""

import json
from io import BytesIO

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from ia import call_claude

MAX_TOKENS_ETAPA5 = 8000


# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------

SYSTEM_COMPARACAO = """
Você é um especialista sênior em procurement produzindo a SÍNTESE TÉCNICA de
um estudo de propostas. Você NÃO está analisando documentos brutos — está
consolidando análises que já foram feitas (baseline, edital, propostas
técnicas avaliadas individualmente). Foco exclusivo em ÓTICA TÉCNICA, escopo
e conformidade. NÃO mencione preço, custo ou valor monetário em nenhum campo.

Você recebe um resumo do baseline, dos requisitos do edital (se houver) e da
avaliação técnica de cada proposta (conformidade por requisito, desvios,
inclusões/exclusões de escopo, flag de não-conformidade mandatória).

Responda SOMENTE com um objeto JSON válido, sem texto antes ou depois:

{
  "n_fornecedores": <número de propostas avaliadas>,
  "eh_comparacao_real": <true se houver 2+ fornecedores, false se houver só 1>,
  "resumo_executivo": "<3-5 frases: o que este one-pager mostra. Se eh_comparacao_real=false, deixe explícito que é leitura de conformidade de um único fornecedor, não uma comparação>",
  "matriz_requisitos": [
    {
      "req_id": "<id do requisito ou '—'>",
      "descricao_curta": "<descrição curta>",
      "tipo": "mandatório|desejável|—",
      "peso": "Alto|Médio|Baixo|—",
      "status_por_fornecedor": [
        {"fornecedor": "<nome>", "status": "cumpre|não cumpre|parcial|desvio"}
      ]
    }
  ],
  "gaps_mandatorios": [
    {
      "fornecedor": "<nome>",
      "requisitos_nao_cumpridos": ["<descrição curta>"],
      "leitura": "<1 frase: severidade do gap nesta ótica técnica, sem falar de preço>"
    }
  ],
  "escopo_cruzado": {
    "inclusoes_exclusivas": [
      {"fornecedor": "<nome>", "item": "<o que só este fornecedor inclui>"}
    ],
    "exclusoes_relevantes": [
      {"fornecedor": "<nome>", "item": "<o que este fornecedor deixa de fora que os outros cobrem ou que o edital pede>"}
    ]
  },
  "leitura_para_decisao": "<parágrafo: o que a ótica técnica, isoladamente, sugere para a decisão — sem indicar 'qual escolher', sem falar de preço. Aponte trade-offs técnicos>",
  "premissas_registradas": ["<premissa assumida nesta síntese>"],
  "faltantes": ["<limitação desta comparação, ex.: só 1 proposta, sem edital, etc.>"]
}

Regras:
- Se houver edital com requisitos, a matriz_requisitos usa os req_id do
  edital. Se não houver edital, construa a matriz a partir dos elementos de
  escopo que aparecem nas avaliações de proposta (use req_id "—").
- status_por_fornecedor: inclua TODOS os fornecedores recebidos em cada linha
  da matriz, mesmo que o fornecedor não tenha avaliado aquele requisito
  explicitamente (nesse caso, status "—" não é uma opção válida — infira a
  partir do escopo declarado ou registre em faltantes que a comparação
  direta não foi possível para aquele item).
- Com 1 fornecedor só: eh_comparacao_real=false. Ainda assim preencha a
  matriz_requisitos (conformidade desse único fornecedor) e gaps_mandatorios
  — é uma leitura útil, só não é comparação. Não invente um segundo
  fornecedor nem finja ranking.
- NUNCA mencione preço, custo, valor, R$, economia. Isso é Etapa 6/7.
- LIMITE DE TAMANHO: matriz_requisitos máximo 30 itens; gaps_mandatorios
  máximo 15; escopo_cruzado.inclusoes_exclusivas e exclusoes_relevantes
  máximo 12 itens cada; premissas e faltantes máximo 6 cada.
  O JSON completo deve caber em 7000 tokens.
"""


# ---------------------------------------------------------------------------
# Montagem de contexto (lê do Estudo, não relê documentos brutos)
# ---------------------------------------------------------------------------

def _resumo_baseline_ctx(estudo) -> str:
    if not estudo.baseline:
        return "Baseline: não disponível."
    tec = estudo.baseline.get("tecnica", {})
    return (
        f"Baseline disponível.\n"
        f"Escopo atual: {tec.get('escopo_atual', '—')}\n"
        f"Fornecedores atuais: {', '.join(tec.get('fornecedores_atuais', []) or ['—'])}"
    )


def _resumo_edital_ctx(estudo) -> str:
    requisitos = (estudo.edital or {}).get("requisitos", [])
    if not requisitos:
        return "Edital: sem requisitos formais disponíveis."
    linhas = ["REQUISITOS DO EDITAL:"]
    for r in requisitos:
        linhas.append(
            f"- [{r.get('id','?')}] ({r.get('tipo','?')}, peso {r.get('peso','?')}) "
            f"{r.get('descricao','')}"
        )
    return "\n".join(linhas)


def _resumo_propostas_ctx(estudo) -> str:
    propostas = estudo.propostas_tecnicas or []
    if not propostas:
        return "Propostas técnicas: nenhuma avaliação disponível."
    blocos = []
    for p in propostas:
        linhas = [f"=== PROPOSTA: {p.get('fornecedor','?')} ==="]
        linhas.append(f"Resumo técnico: {p.get('resumo_tecnico','—')}")
        conf = p.get("conformidade", [])
        if conf:
            linhas.append("Conformidade por requisito:")
            for c in conf:
                linhas.append(
                    f"  [{c.get('req_id','—')}] {c.get('descricao_curta','')} "
                    f"({c.get('tipo','—')}) -> {c.get('status','—')} "
                    f"— {c.get('observacao','')}"
                )
        if p.get("nao_cumpre_mandatorio"):
            linhas.append(
                f"Mandatórios não cumpridos: {', '.join(p.get('mandatorios_nao_cumpridos', []))}"
            )
        incl = p.get("inclusoes_escopo", [])
        if incl:
            linhas.append(f"Inclusões de escopo: {', '.join(incl)}")
        excl = p.get("exclusoes_escopo", [])
        if excl:
            linhas.append(f"Exclusões de escopo: {', '.join(excl)}")
        blocos.append("\n".join(linhas))
    return "\n\n".join(blocos)


def _parse_resposta(resposta_bruta: str) -> dict:
    """Parser com fallback de truncamento (mesmo padrão das outras etapas)."""
    texto = resposta_bruta.strip()
    if texto.startswith("```"):
        linhas = texto.split("\n")
        texto = "\n".join(linhas[1:-1]).strip()
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        texto_cortado = texto.rstrip().rstrip(",")
        abertas = texto_cortado.count("{") - texto_cortado.count("}")
        colchetes = texto_cortado.count("[") - texto_cortado.count("]")
        fechamento = "]" * colchetes + "}" * abertas
        try:
            return json.loads(texto_cortado + fechamento)
        except json.JSONDecodeError:
            raise ValueError(
                "JSON inválido mesmo após tentativa de correção. "
                "Estudo com muitos fornecedores/requisitos — verifique o volume de dados."
            )


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def rodar_etapa5(estudo) -> dict:
    """
    Executa a Etapa 5 completa.

    Retorna dict com:
        - tem_dados : bool
        - analise   : dict completo (JSON do Claude) ou None
        - resumo    : texto legível pra UI
    """
    if not estudo.propostas_tecnicas:
        msg = "Nenhuma proposta técnica avaliada — Etapa 5 não pôde rodar."
        estudo.add_faltante(msg)
        estudo.comparacao_tecnica = None
        return {"tem_dados": False, "analise": None, "resumo": f"⚠️ {msg}"}

    contexto = (
        f"Categoria: {estudo.categoria}\n"
        f"Modelo de precificação: {estudo.modelo_precificacao}\n"
        f"Micro-categoria: {estudo.micro_categoria or '—'}\n\n"
        f"{_resumo_baseline_ctx(estudo)}\n\n"
        f"{_resumo_edital_ctx(estudo)}\n\n"
        f"{_resumo_propostas_ctx(estudo)}"
    )

    with st.spinner("Etapa 5 — consolidando comparação técnica..."):
        resposta_bruta = call_claude(
            messages=[{"role": "user", "content": contexto}],
            system=SYSTEM_COMPARACAO,
            max_tokens=MAX_TOKENS_ETAPA5,
        )

    try:
        analise = _parse_resposta(resposta_bruta)
    except (json.JSONDecodeError, ValueError) as e:
        st.error(f"Erro ao interpretar resposta da IA na Etapa 5: {e}\n\nResposta bruta:\n{resposta_bruta}")
        st.stop()

    estudo.comparacao_tecnica = analise

    for p in analise.get("premissas_registradas", []):
        estudo.add_premissa(p)
    for f in analise.get("faltantes", []):
        estudo.add_faltante(f)
    if not analise.get("eh_comparacao_real", True):
        estudo.add_faltante(
            "Etapa 5 rodou com apenas 1 proposta técnica — não há comparação "
            "entre fornecedores, só leitura de conformidade individual."
        )

    estudo.etapa_atual = 5

    resumo = _montar_resumo(analise)
    return {"tem_dados": True, "analise": analise, "resumo": resumo}


def _montar_resumo(analise: dict) -> str:
    linhas = []

    resumo_exec = analise.get("resumo_executivo")
    if resumo_exec:
        linhas.append(f"**Resumo executivo:** {resumo_exec}")

    if not analise.get("eh_comparacao_real", True):
        linhas.append("\n⚠️ _Apenas 1 fornecedor — leitura de conformidade individual, não comparação._")

    matriz = analise.get("matriz_requisitos", [])
    if matriz:
        linhas.append(f"\n**Matriz de requisitos:** {len(matriz)} itens consolidados.")

    gaps = analise.get("gaps_mandatorios", [])
    if gaps:
        linhas.append("\n**Gaps mandatórios por fornecedor:**")
        for g in gaps:
            reqs = ", ".join(g.get("requisitos_nao_cumpridos", [])) or "(ver detalhe)"
            linhas.append(f"- **{g.get('fornecedor','?')}**: {reqs} — {g.get('leitura','')}")

    escopo = analise.get("escopo_cruzado", {})
    inclusoes = escopo.get("inclusoes_exclusivas", [])
    if inclusoes:
        linhas.append("\n**Inclusões exclusivas de escopo:**")
        for i in inclusoes:
            linhas.append(f"- **{i.get('fornecedor','?')}**: {i.get('item','')}")
    exclusoes = escopo.get("exclusoes_relevantes", [])
    if exclusoes:
        linhas.append("\n**Exclusões relevantes:**")
        for e in exclusoes:
            linhas.append(f"- **{e.get('fornecedor','?')}**: {e.get('item','')}")

    leitura = analise.get("leitura_para_decisao")
    if leitura:
        linhas.append(f"\n**Leitura para decisão (ótica técnica, sem preço):**\n{leitura}")

    faltantes = analise.get("faltantes", [])
    if faltantes:
        linhas.append("\n**Limitações desta comparação:**")
        for f in faltantes:
            linhas.append(f"- {f}")

    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Geração Word (one-pager)
# ---------------------------------------------------------------------------

COR_AZUL_AM = RGBColor(0x1F, 0x3A, 0x5F)
COR_CINZA = RGBColor(0x59, 0x59, 0x59)

STATUS_COR = {
    "cumpre": "C6E8C6",
    "parcial": "FFF2B2",
    "não cumpre": "F4C7C3",
    "desvio": "D9D2E9",
}


def gerar_word_etapa5(estudo) -> BytesIO:
    """Gera o one-pager Word da Etapa 5. Retorna BytesIO pronto para download."""
    analise = estudo.comparacao_tecnica or {}
    fornecedores = [p.get("fornecedor", "?") for p in (estudo.propostas_tecnicas or [])]

    doc = Document()

    # Margens enxutas — one-pager
    secao = doc.sections[0]
    secao.left_margin = Cm(1.8)
    secao.right_margin = Cm(1.8)
    secao.top_margin = Cm(1.5)
    secao.bottom_margin = Cm(1.5)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(10)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = titulo.add_run("Comparação Técnica — One-Pager")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM

    sub = doc.add_paragraph()
    run_sub = sub.add_run(
        f"{estudo.micro_categoria or estudo.categoria or 'Categoria não identificada'} · "
        f"{len(fornecedores)} fornecedor(es) avaliado(s) · ótica técnica, sem preço"
    )
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = COR_CINZA
    run_sub.italic = True

    _linha_divisoria(doc)

    if not analise.get("eh_comparacao_real", True):
        aviso = doc.add_paragraph()
        run_aviso = aviso.add_run(
            "⚠ Apenas 1 fornecedor avaliado — este documento mostra a conformidade "
            "individual, não uma comparação entre propostas."
        )
        run_aviso.font.size = Pt(9)
        run_aviso.italic = True
        run_aviso.font.color.rgb = COR_CINZA

    # Resumo executivo
    if analise.get("resumo_executivo"):
        _heading(doc, "Resumo Executivo")
        doc.add_paragraph(analise["resumo_executivo"])

    # Matriz de requisitos
    matriz = analise.get("matriz_requisitos", [])
    if matriz:
        _heading(doc, "Matriz de Conformidade")
        _tabela_matriz_word(doc, matriz, fornecedores)

    # Gaps mandatórios
    gaps = analise.get("gaps_mandatorios", [])
    if gaps:
        _heading(doc, "Gaps Mandatórios")
        for g in gaps:
            p = doc.add_paragraph(style=None)
            run_f = p.add_run(f"{g.get('fornecedor','?')}: ")
            run_f.bold = True
            run_f.font.size = Pt(10)
            reqs = ", ".join(g.get("requisitos_nao_cumpridos", [])) or "—"
            p.add_run(f"{reqs} — {g.get('leitura','')}").font.size = Pt(10)

    # Escopo cruzado
    escopo = analise.get("escopo_cruzado", {})
    inclusoes = escopo.get("inclusoes_exclusivas", [])
    exclusoes = escopo.get("exclusoes_relevantes", [])
    if inclusoes or exclusoes:
        _heading(doc, "Escopo — Diferenças entre Fornecedores")
        if inclusoes:
            p = doc.add_paragraph()
            p.add_run("Inclusões exclusivas:").bold = True
            for i in inclusoes:
                doc.add_paragraph(f"{i.get('fornecedor','?')}: {i.get('item','')}", style="List Bullet")
        if exclusoes:
            p = doc.add_paragraph()
            p.add_run("Exclusões relevantes:").bold = True
            for e in exclusoes:
                doc.add_paragraph(f"{e.get('fornecedor','?')}: {e.get('item','')}", style="List Bullet")

    # Leitura para decisão
    if analise.get("leitura_para_decisao"):
        _heading(doc, "Leitura para Decisão (ótica técnica)")
        doc.add_paragraph(analise["leitura_para_decisao"])

    # Limitações
    faltantes = analise.get("faltantes", [])
    if faltantes:
        _heading(doc, "Limitações desta Comparação")
        for f in faltantes:
            doc.add_paragraph(f, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _heading(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM


def _linha_divisoria(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _tabela_matriz_word(doc, matriz, fornecedores):
    n_cols = 3 + len(fornecedores)  # req_id, descrição, tipo, + 1 col por fornecedor
    tabela = doc.add_table(rows=1, cols=n_cols)
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabela.style = "Table Grid"

    cabecalho = ["Req.", "Descrição", "Tipo"] + fornecedores
    hdr_cells = tabela.rows[0].cells
    for i, texto in enumerate(cabecalho):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(texto)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _sombrear_celula(hdr_cells[i], "1F3A5F")

    for item in matriz:
        row = tabela.add_row().cells
        row[0].text = ""
        row[0].paragraphs[0].add_run(item.get("req_id", "—")).font.size = Pt(8)
        row[1].text = ""
        row[1].paragraphs[0].add_run(item.get("descricao_curta", "")).font.size = Pt(8)
        row[2].text = ""
        row[2].paragraphs[0].add_run(item.get("tipo", "—")).font.size = Pt(8)

        status_map = {
            s.get("fornecedor"): s.get("status", "—")
            for s in item.get("status_por_fornecedor", [])
        }
        for i, forn in enumerate(fornecedores):
            status = status_map.get(forn, "—")
            cell = row[3 + i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(status)
            run.font.size = Pt(8)
            cor_hex = STATUS_COR.get(status)
            if cor_hex:
                _sombrear_celula(cell, cor_hex)

    # Larguras
    larguras = [Cm(1.3), Cm(6.0), Cm(1.8)] + [Cm(2.5)] * len(fornecedores)
    for row in tabela.rows:
        for i, cell in enumerate(row.cells):
            cell.width = larguras[i]


def _sombrear_celula(cell, cor_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Geração Excel (matriz detalhada)
# ---------------------------------------------------------------------------

FILL_HEADER = PatternFill(start_color="1F3A5F", end_color="1F3A5F", fill_type="solid")
FONT_HEADER = Font(name="Arial", size=10, bold=True, color="FFFFFF")
FONT_BODY = Font(name="Arial", size=10)
FONT_BODY_BOLD = Font(name="Arial", size=10, bold=True)
THIN = Side(style="thin", color="CCCCCC")
BORDA = Border(top=THIN, bottom=THIN, left=THIN, right=THIN)

STATUS_FILL = {
    "cumpre": PatternFill(start_color="C6E8C6", end_color="C6E8C6", fill_type="solid"),
    "parcial": PatternFill(start_color="FFF2B2", end_color="FFF2B2", fill_type="solid"),
    "não cumpre": PatternFill(start_color="F4C7C3", end_color="F4C7C3", fill_type="solid"),
    "desvio": PatternFill(start_color="D9D2E9", end_color="D9D2E9", fill_type="solid"),
}


def gerar_excel_etapa5(estudo) -> BytesIO:
    """Gera o Excel da Etapa 5 (3 abas). Retorna BytesIO pronto para download."""
    analise = estudo.comparacao_tecnica or {}
    fornecedores = [p.get("fornecedor", "?") for p in (estudo.propostas_tecnicas or [])]

    wb = Workbook()

    _aba_matriz(wb, analise, fornecedores)
    _aba_gaps(wb, analise)
    _aba_escopo(wb, analise)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _cabecalho_aba(ws, colunas):
    for i, texto in enumerate(colunas, start=1):
        cell = ws.cell(row=1, column=i, value=texto)
        cell.font = FONT_HEADER
        cell.fill = FILL_HEADER
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = BORDA
    ws.row_dimensions[1].height = 28
    ws.freeze_panes = "A2"


def _aba_matriz(wb, analise, fornecedores):
    ws = wb.active
    ws.title = "Matriz de Conformidade"

    colunas = ["Req.", "Descrição", "Tipo", "Peso"] + fornecedores
    _cabecalho_aba(ws, colunas)

    matriz = analise.get("matriz_requisitos", [])
    for r, item in enumerate(matriz, start=2):
        ws.cell(row=r, column=1, value=item.get("req_id", "—")).font = FONT_BODY
        ws.cell(row=r, column=2, value=item.get("descricao_curta", "")).font = FONT_BODY
        ws.cell(row=r, column=3, value=item.get("tipo", "—")).font = FONT_BODY
        ws.cell(row=r, column=4, value=item.get("peso", "—")).font = FONT_BODY

        status_map = {
            s.get("fornecedor"): s.get("status", "—")
            for s in item.get("status_por_fornecedor", [])
        }
        for i, forn in enumerate(fornecedores):
            status = status_map.get(forn, "—")
            cell = ws.cell(row=r, column=5 + i, value=status)
            cell.font = FONT_BODY
            cell.alignment = Alignment(horizontal="center")
            fill = STATUS_FILL.get(status)
            if fill:
                cell.fill = fill

        for c in range(1, len(colunas) + 1):
            ws.cell(row=r, column=c).border = BORDA

    larguras = [8, 45, 12, 10] + [18] * len(fornecedores)
    for i, larg in enumerate(larguras, start=1):
        ws.column_dimensions[get_column_letter(i)].width = larg


def _aba_gaps(wb, analise):
    ws = wb.create_sheet("Gaps Mandatórios")
    _cabecalho_aba(ws, ["Fornecedor", "Requisitos não cumpridos", "Leitura técnica"])

    gaps = analise.get("gaps_mandatorios", [])
    for r, g in enumerate(gaps, start=2):
        ws.cell(row=r, column=1, value=g.get("fornecedor", "?")).font = FONT_BODY_BOLD
        ws.cell(row=r, column=2, value=", ".join(g.get("requisitos_nao_cumpridos", []))).font = FONT_BODY
        ws.cell(row=r, column=3, value=g.get("leitura", "")).font = FONT_BODY
        for c in range(1, 4):
            cell = ws.cell(row=r, column=c)
            cell.border = BORDA
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 50


def _bloco_secao_escopo(ws, r_inicio, titulo, itens):
    """Escreve um bloco de seção (título + cabeçalho + linhas) na aba de escopo.
    Retorna a próxima linha livre após o bloco."""
    ws.cell(row=r_inicio, column=1, value=titulo).font = FONT_HEADER
    for c in range(1, 3):
        ws.cell(row=r_inicio, column=c).fill = FILL_HEADER
    ws.merge_cells(start_row=r_inicio, start_column=1, end_row=r_inicio, end_column=2)
    ws.row_dimensions[r_inicio].height = 20

    r = r_inicio + 1
    ws.cell(row=r, column=1, value="Fornecedor").font = FONT_BODY_BOLD
    ws.cell(row=r, column=2, value="Item").font = FONT_BODY_BOLD
    for c in range(1, 3):
        ws.cell(row=r, column=c).fill = PatternFill(start_color="DCE6F0", end_color="DCE6F0", fill_type="solid")
    r += 1

    if not itens:
        ws.cell(row=r, column=1, value="—").font = FONT_BODY
        ws.cell(row=r, column=2, value="(nenhum registrado)").font = FONT_BODY
        r += 1
    else:
        for item in itens:
            ws.cell(row=r, column=1, value=item.get("fornecedor", "?")).font = FONT_BODY
            ws.cell(row=r, column=2, value=item.get("item", "")).font = FONT_BODY
            r += 1

    return r + 1  # linha em branco de separação


def _aba_escopo(wb, analise):
    ws = wb.create_sheet("Escopo — Diferenças")
    escopo = analise.get("escopo_cruzado", {})
    inclusoes = escopo.get("inclusoes_exclusivas", [])
    exclusoes = escopo.get("exclusoes_relevantes", [])

    proxima_linha = _bloco_secao_escopo(ws, 1, "Inclusões exclusivas", inclusoes)
    _bloco_secao_escopo(ws, proxima_linha, "Exclusões relevantes", exclusoes)

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=2):
        for cell in row:
            if cell.value is not None:
                cell.border = BORDA
                cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 60

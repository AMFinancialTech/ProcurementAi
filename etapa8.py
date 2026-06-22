"""
etapa8.py — Etapa 8: Estratégia da Categoria (Matriz de Kraljic).

Módulo STANDALONE — diferente das Etapas 2-7, não depende da cadeia de
análise de propostas. Pode rodar:
- Aproveitando o Estudo: se as Etapas 1-2 já rodaram, puxa categoria e gasto
  anual do baseline.
- Do zero: o consultor informa a categoria manualmente.

O que faz:
1. Classifica a categoria de compra na Matriz de Kraljic, dois eixos:
   - Impacto financeiro/no resultado (do gasto do baseline, ou perguntado).
   - Risco de suprimento (TENTA via web search; o consultor pode ajustar).
2. Posiciona num dos 4 quadrantes (alavancagem, estratégico, não-crítico,
   gargalo) e gera a recomendação estratégica do quadrante.
3. Monta a árvore de categorias (macro → categoria → subcategoria).
4. Grava em estudo.estrategia_categoria.
5. Gera Word (one-pager de estratégia) e Excel. Mesmo estilo visual das
   Etapas 5 e 7. PPT: pendente layout (mesma regra das outras etapas).

Sobre o risco de suprimento (decisão do usuário: "os dois"):
- rodar_etapa8 TENTA pesquisar o risco via web search e grava como SUGESTÃO.
- O consultor pode sobrescrever via chat (confirmar_risco_suprimento).
- Se a busca falhar (ferramenta indisponível, erro), cai no modo manual sem
  travar — registra faltante e segue com o que o consultor informar.
"""

import json
from io import BytesIO

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from ia import call_claude, call_claude_com_busca

MAX_TOKENS_ETAPA8 = 8000

# Os 4 quadrantes da Matriz de Kraljic, indexados por (impacto, risco).
QUADRANTES = {
    ("alto", "alto"): "estrategico",
    ("alto", "baixo"): "alavancagem",
    ("baixo", "alto"): "gargalo",
    ("baixo", "baixo"): "nao_critico",
}

NOME_QUADRANTE = {
    "estrategico": "Estratégico",
    "alavancagem": "Alavancagem",
    "gargalo": "Gargalo",
    "nao_critico": "Não-crítico",
}


# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

SYSTEM_RISCO_SUPRIMENTO = """
Você é um especialista sênior em procurement avaliando o RISCO DE SUPRIMENTO
de uma categoria de compra, usando a Matriz de Kraljic como referência.

Use a ferramenta de busca para pesquisar o mercado fornecedor da categoria
informada: número de fornecedores disponíveis no mercado, concentração de
mercado, dependência geográfica, escassez de insumos, volatilidade de preço,
barreiras de entrada, e qualquer fator que afete a facilidade/dificuldade de
garantir o fornecimento. Foque no mercado brasileiro quando aplicável.

Responda SOMENTE com um objeto JSON válido, sem texto antes ou depois:

{
  "risco_suprimento": "alto|baixo",
  "justificativa": "<2-4 frases explicando a classificação, citando os fatores de mercado encontrados>",
  "fatores_risco": ["<fator concreto que aumenta ou reduz o risco>"],
  "fontes_consultadas": ["<o que foi pesquisado, em termos gerais — não precisa de URL>"],
  "confianca": "alta|média|baixa",
  "premissas": ["<premissa assumida>"],
  "faltantes": ["<o que não foi possível determinar pela pesquisa>"]
}

Regras:
- risco_suprimento é BINÁRIO (alto ou baixo) — é o eixo Y da Kraljic. Mesmo
  com nuance, escolha o lado predominante e explique na justificativa.
- Se a pesquisa não trouxer dados suficientes, use confianca "baixa" e
  registre em faltantes — não invente dados de mercado específicos.
- LIMITE DE TAMANHO: fatores_risco máx. 6, fontes_consultadas máx. 6,
  premissas e faltantes máx. 5 cada. O JSON deve caber em 4000 tokens.
"""

SYSTEM_ESTRATEGIA = """
Você é um especialista sênior em procurement produzindo a ESTRATÉGIA DE
CATEGORIA usando a Matriz de Kraljic. Você recebe a categoria, o eixo de
impacto financeiro (alto/baixo), o eixo de risco de suprimento (alto/baixo),
e o quadrante Kraljic resultante. NÃO reclassifique os eixos — eles já foram
decididos. Sua tarefa é produzir a estratégia e a árvore de categorias.

Os 4 quadrantes e suas estratégias clássicas:
- Estratégico (impacto alto + risco alto): poucos fornecedores críticos,
  parceria de longo prazo, gestão de relacionamento, planos de contingência.
- Alavancagem (impacto alto + risco baixo): muitos fornecedores, use poder de
  barganha, concorrência, foco em redução de preço, leilões reversos.
- Gargalo (impacto baixo + risco alto): garanta o fornecimento, contratos de
  longo prazo, estoque de segurança, fornecedores alternativos, reduza
  dependência.
- Não-crítico (impacto baixo + risco baixo): simplifique, automatize compras,
  reduza esforço administrativo, padronize, catálogos.

Responda SOMENTE com um objeto JSON válido, sem texto antes ou depois:

{
  "quadrante": "estrategico|alavancagem|gargalo|nao_critico",
  "resumo_posicao": "<2-3 frases: por que a categoria está neste quadrante, dado impacto e risco>",
  "estrategia_recomendada": "<parágrafo: a abordagem estratégica recomendada para esta categoria, adaptada ao contexto específico — não genérica>",
  "acoes_taticas": [
    {"acao": "<ação concreta>", "prazo": "curto|médio|longo", "racional": "<1 frase>"}
  ],
  "numero_fornecedores_sugerido": "<sugestão de quantos fornecedores faz sentido para esta categoria e por quê>",
  "tipo_relacionamento": "<transacional|colaborativo|parceria estratégica — e por quê>",
  "arvore_categoria": {
    "macro": "<macro-categoria, ex.: Serviços, Materiais, Commodities>",
    "categoria": "<a categoria analisada>",
    "subcategorias": ["<subcategoria relevante>"]
  },
  "alertas_estrategicos": ["<risco ou ponto de atenção estratégico de longo prazo>"],
  "premissas": ["<premissa assumida>"],
  "faltantes": ["<limitação desta análise estratégica>"]
}

Regras:
- O campo quadrante DEVE bater com o que foi informado — não recalcule.
- estrategia_recomendada e acoes_taticas devem ser específicas para a
  categoria, não cópia genérica da definição do quadrante.
- LIMITE DE TAMANHO: acoes_taticas máx. 8, subcategorias máx. 10,
  alertas_estrategicos máx. 6, premissas e faltantes máx. 5 cada.
  O JSON completo deve caber em 7000 tokens.
"""


# ---------------------------------------------------------------------------
# Resolução de contexto (standalone OU aproveitando o Estudo)
# ---------------------------------------------------------------------------

def _resolver_categoria(estudo, categoria_manual: str | None) -> str:
    """Decide a categoria: manual (se informada) tem prioridade; senão, do Estudo."""
    if categoria_manual:
        return categoria_manual
    if estudo.micro_categoria:
        return estudo.micro_categoria
    if estudo.categoria:
        return estudo.categoria
    return ""


def _resolver_impacto(estudo, impacto_manual: str | None) -> tuple[str, str]:
    """
    Decide o eixo de impacto financeiro (alto/baixo).
    Retorna (impacto, origem_do_dado).
    Manual tem prioridade. Senão, infere do gasto anual do baseline.
    """
    if impacto_manual in ("alto", "baixo"):
        return impacto_manual, "informado pelo consultor"

    com = (estudo.baseline or {}).get("comercial", {})
    gasto = com.get("preco_anual_total")
    if isinstance(gasto, (int, float)):
        # Heurística simples e transparente: acima de R$ 1 mi/ano = impacto alto.
        # O consultor pode sempre sobrescrever via chat.
        if gasto >= 1_000_000:
            return "alto", f"inferido do gasto anual do baseline (R$ {gasto:,.0f})"
        return "baixo", f"inferido do gasto anual do baseline (R$ {gasto:,.0f})"

    return "", "não disponível"


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

def _parse_resposta(resposta_bruta: str) -> dict:
    """Parser com fallback de truncamento (mesmo padrão das outras etapas)."""
    texto = resposta_bruta.strip()
    if texto.startswith("```"):
        linhas = texto.split("\n")
        texto = "\n".join(linhas[1:-1]).strip()
    # A resposta com busca pode ter texto antes do JSON — tenta achar o objeto.
    if not texto.startswith("{"):
        inicio = texto.find("{")
        if inicio != -1:
            texto = texto[inicio:]
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        texto_cortado = texto.rstrip().rstrip(",")
        abertas = texto_cortado.count("{") - texto_cortado.count("}")
        colchetes = texto_cortado.count("[") - texto_cortado.count("]")
        fechamento = "]" * colchetes + "}" * abertas
        try:
            return json.loads(texto_cortado + fechamento)
        except json.JSONDecodeError:
            raise ValueError(
                "JSON inválido mesmo após tentativa de correção na Etapa 8."
            )


# ---------------------------------------------------------------------------
# Risco de suprimento: pesquisa (com fallback) e confirmação manual
# ---------------------------------------------------------------------------

def pesquisar_risco_suprimento(categoria: str) -> dict:
    """
    Tenta pesquisar o risco de suprimento via web search.

    Retorna dict com:
        - sucesso       : bool
        - risco         : "alto"|"baixo"|None
        - analise       : dict completo da IA (ou None)
        - houve_busca   : bool
        - erro          : str (se sucesso=False)
    """
    contexto = (
        f"Categoria de compra a avaliar: {categoria}\n\n"
        f"Pesquise o mercado fornecedor desta categoria e classifique o risco "
        f"de suprimento (eixo Y da Matriz de Kraljic)."
    )
    try:
        resposta_bruta, houve_busca = call_claude_com_busca(
            messages=[{"role": "user", "content": contexto}],
            system=SYSTEM_RISCO_SUPRIMENTO,
            max_tokens=4000,
            max_buscas=3,
        )
        analise = _parse_resposta(resposta_bruta)
        return {
            "sucesso": True,
            "risco": analise.get("risco_suprimento"),
            "analise": analise,
            "houve_busca": houve_busca,
            "erro": None,
        }
    except Exception as e:  # noqa: BLE001 — qualquer falha cai no modo manual
        return {
            "sucesso": False,
            "risco": None,
            "analise": None,
            "houve_busca": False,
            "erro": str(e),
        }


def confirmar_risco_suprimento(estudo, risco: str) -> str:
    """
    Registra o risco de suprimento informado/confirmado pelo consultor.
    Não chama IA — atribuição direta (mesmo padrão da taxa/moeda da Etapa 6).
    Guarda num campo temporário do Estudo até rodar_etapa8 montar a estratégia.
    """
    risco = risco.strip().lower()
    if risco not in ("alto", "baixo"):
        return "Risco inválido. Informe 'alto' ou 'baixo'."
    # Usa um atributo dinâmico simples para carregar a decisão até a montagem.
    estudo._risco_suprimento_confirmado = risco
    estudo.add_premissa(f"Risco de suprimento da categoria definido como '{risco}'.")
    return f"✅ Risco de suprimento registrado como **{risco}**."


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def rodar_etapa8(estudo, categoria_manual=None, impacto_manual=None, risco_manual=None) -> dict:
    """
    Executa a Etapa 8 (Estratégia da Categoria).

    Parâmetros opcionais (para o modo standalone / sobrescrita do consultor):
        - categoria_manual : str — categoria, se as Etapas 1-2 não rodaram
        - impacto_manual   : "alto"|"baixo" — sobrescreve a inferência de impacto
        - risco_manual     : "alto"|"baixo" — sobrescreve a pesquisa de risco

    Retorna dict com:
        - tem_dados : bool
        - analise   : dict completo ou None
        - resumo    : texto legível
        - precisa_categoria : bool (True se não há categoria nem do Estudo nem manual)
    """
    categoria = _resolver_categoria(estudo, categoria_manual)
    if not categoria:
        return {
            "tem_dados": False,
            "analise": None,
            "resumo": "⚠️ Nenhuma categoria identificada. Informe a categoria de compra para rodar a Etapa 8.",
            "precisa_categoria": True,
        }

    # --- Eixo de impacto financeiro ---
    impacto, origem_impacto = _resolver_impacto(estudo, impacto_manual)

    # --- Eixo de risco de suprimento ---
    risco = None
    analise_risco = None
    origem_risco = ""

    if risco_manual in ("alto", "baixo"):
        risco = risco_manual
        origem_risco = "informado pelo consultor"
    elif getattr(estudo, "_risco_suprimento_confirmado", None) in ("alto", "baixo"):
        risco = estudo._risco_suprimento_confirmado
        origem_risco = "confirmado pelo consultor"
    else:
        with st.spinner("Etapa 8 — pesquisando risco de suprimento da categoria..."):
            resultado_busca = pesquisar_risco_suprimento(categoria)
        if resultado_busca["sucesso"] and resultado_busca["risco"] in ("alto", "baixo"):
            risco = resultado_busca["risco"]
            analise_risco = resultado_busca["analise"]
            origem_risco = (
                "pesquisado via web search" if resultado_busca["houve_busca"]
                else "estimado pela IA (sem busca efetiva)"
            )
            for p in (analise_risco.get("premissas", []) if analise_risco else []):
                estudo.add_premissa(f"[risco suprimento] {p}")
            for f in (analise_risco.get("faltantes", []) if analise_risco else []):
                estudo.add_faltante(f"[risco suprimento] {f}")
        else:
            erro = resultado_busca.get("erro", "motivo desconhecido")
            estudo.add_faltante(
                f"Pesquisa de risco de suprimento falhou ({erro}). "
                f"Risco precisa ser informado manualmente."
            )

    # Se ainda não temos impacto ou risco, sinaliza que precisa de input humano.
    if impacto not in ("alto", "baixo") or risco not in ("alto", "baixo"):
        return {
            "tem_dados": False,
            "analise": None,
            "resumo": (
                f"⚠️ Para montar a Matriz de Kraljic da categoria **{categoria}**, "
                f"ainda preciso definir:\n"
                + ("" if impacto in ("alto", "baixo") else "- **Impacto financeiro** (alto/baixo)\n")
                + ("" if risco in ("alto", "baixo") else "- **Risco de suprimento** (alto/baixo)\n")
            ),
            "precisa_categoria": False,
            "falta_impacto": impacto not in ("alto", "baixo"),
            "falta_risco": risco not in ("alto", "baixo"),
            "categoria": categoria,
            "analise_risco": analise_risco,
        }

    # --- Quadrante Kraljic ---
    quadrante = QUADRANTES[(impacto, risco)]

    # --- Estratégia (chamada de IA, sem busca) ---
    contexto = (
        f"Categoria: {categoria}\n"
        f"Macro-categoria (se conhecida): {estudo.categoria or '—'}\n"
        f"Impacto financeiro: {impacto} ({origem_impacto})\n"
        f"Risco de suprimento: {risco} ({origem_risco})\n"
        f"Quadrante Kraljic resultante: {NOME_QUADRANTE[quadrante]}\n"
    )
    if analise_risco:
        contexto += f"\nJustificativa do risco (da pesquisa): {analise_risco.get('justificativa', '—')}"

    with st.spinner("Etapa 8 — montando estratégia da categoria..."):
        resposta_bruta = call_claude(
            messages=[{"role": "user", "content": contexto}],
            system=SYSTEM_ESTRATEGIA,
            max_tokens=MAX_TOKENS_ETAPA8,
        )

    try:
        analise = _parse_resposta(resposta_bruta)
    except (json.JSONDecodeError, ValueError) as e:
        st.error(f"Erro ao interpretar resposta da IA na Etapa 8: {e}\n\nResposta bruta:\n{resposta_bruta}")
        st.stop()

    # Anexa os dados dos eixos e da pesquisa ao resultado, para a saída formatada.
    analise["_impacto"] = impacto
    analise["_origem_impacto"] = origem_impacto
    analise["_risco"] = risco
    analise["_origem_risco"] = origem_risco
    analise["_categoria"] = categoria
    analise["_analise_risco"] = analise_risco

    estudo.estrategia_categoria = analise

    for p in analise.get("premissas", []):
        estudo.add_premissa(p)
    for f in analise.get("faltantes", []):
        estudo.add_faltante(f)

    estudo.etapa_atual = 8

    resumo = _montar_resumo(analise)
    return {"tem_dados": True, "analise": analise, "resumo": resumo, "precisa_categoria": False}


def _montar_resumo(analise: dict) -> str:
    linhas = []

    quadrante = analise.get("quadrante", "")
    nome_q = NOME_QUADRANTE.get(quadrante, quadrante)
    linhas.append(
        f"## Estratégia da Categoria: {analise.get('_categoria', '—')}\n"
    )
    linhas.append(
        f"**Quadrante Kraljic: {nome_q}** "
        f"(impacto {analise.get('_impacto', '—')} · risco de suprimento {analise.get('_risco', '—')})"
    )
    linhas.append(f"_Impacto: {analise.get('_origem_impacto', '—')}_")
    linhas.append(f"_Risco: {analise.get('_origem_risco', '—')}_")

    if analise.get("resumo_posicao"):
        linhas.append(f"\n**Posição:** {analise['resumo_posicao']}")

    if analise.get("estrategia_recomendada"):
        linhas.append(f"\n**Estratégia recomendada:**\n{analise['estrategia_recomendada']}")

    acoes = analise.get("acoes_taticas", [])
    if acoes:
        linhas.append("\n**Ações táticas:**")
        for a in acoes:
            linhas.append(f"- [{a.get('prazo', '—')}] {a.get('acao', '')} — {a.get('racional', '')}")

    if analise.get("numero_fornecedores_sugerido"):
        linhas.append(f"\n**Número de fornecedores sugerido:** {analise['numero_fornecedores_sugerido']}")
    if analise.get("tipo_relacionamento"):
        linhas.append(f"**Tipo de relacionamento:** {analise['tipo_relacionamento']}")

    arvore = analise.get("arvore_categoria", {})
    if arvore:
        subs = ", ".join(arvore.get("subcategorias", [])) or "—"
        linhas.append(
            f"\n**Árvore de categoria:** {arvore.get('macro', '—')} → "
            f"{arvore.get('categoria', '—')} → [{subs}]"
        )

    alertas = analise.get("alertas_estrategicos", [])
    if alertas:
        linhas.append("\n**Alertas estratégicos:**")
        for al in alertas:
            linhas.append(f"- {al}")

    faltantes = analise.get("faltantes", [])
    if faltantes:
        linhas.append("\n**Limitações desta análise:**")
        for f in faltantes:
            linhas.append(f"- {f}")

    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Geração Word — mesmo estilo visual das Etapas 5 e 7
# ---------------------------------------------------------------------------

COR_AZUL_AM = RGBColor(0x1F, 0x3A, 0x5F)
COR_CINZA = RGBColor(0x59, 0x59, 0x59)

# Cor de fundo de cada quadrante (para destacar o quadrante ativo na matriz visual).
QUADRANTE_COR = {
    "estrategico": "F4C7C3",   # vermelho claro — atenção máxima
    "alavancagem": "C6E8C6",   # verde claro — oportunidade
    "gargalo": "FFF2B2",       # amarelo — cuidado
    "nao_critico": "DCE6F0",   # azul claro — rotina
}


def _heading(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM


def _linha_divisoria(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _sombrear_celula(cell, cor_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


def _matriz_kraljic_word(doc, quadrante_ativo):
    """Desenha a matriz 2x2 de Kraljic, destacando o quadrante ativo."""
    # Linhas: impacto alto (cima), impacto baixo (baixo).
    # Colunas: risco baixo (esq), risco alto (dir).
    tabela = doc.add_table(rows=3, cols=3)
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalhos de eixo
    def _set(cell, texto, bold=False, size=8, cor=None, fill=None):
        cell.text = ""
        run = cell.paragraphs[0].add_run(texto)
        run.font.size = Pt(size)
        run.font.bold = bold
        if cor:
            run.font.color.rgb = cor
        if fill:
            _sombrear_celula(cell, fill)

    # Linha 0: canto vazio + dois rótulos de risco
    _set(tabela.rows[0].cells[0], "", size=8)
    _set(tabela.rows[0].cells[1], "Risco de suprimento: BAIXO", bold=True, size=8, cor=RGBColor(0xFF,0xFF,0xFF), fill="1F3A5F")
    _set(tabela.rows[0].cells[2], "Risco de suprimento: ALTO", bold=True, size=8, cor=RGBColor(0xFF,0xFF,0xFF), fill="1F3A5F")

    # Linha 1: rótulo "impacto alto" + Alavancagem + Estratégico
    _set(tabela.rows[1].cells[0], "Impacto\nALTO", bold=True, size=8, cor=RGBColor(0xFF,0xFF,0xFF), fill="1F3A5F")
    _celula_quadrante(tabela.rows[1].cells[1], "alavancagem", quadrante_ativo)
    _celula_quadrante(tabela.rows[1].cells[2], "estrategico", quadrante_ativo)

    # Linha 2: rótulo "impacto baixo" + Não-crítico + Gargalo
    _set(tabela.rows[2].cells[0], "Impacto\nBAIXO", bold=True, size=8, cor=RGBColor(0xFF,0xFF,0xFF), fill="1F3A5F")
    _celula_quadrante(tabela.rows[2].cells[1], "nao_critico", quadrante_ativo)
    _celula_quadrante(tabela.rows[2].cells[2], "gargalo", quadrante_ativo)

    larguras = [Cm(2.5), Cm(5.5), Cm(5.5)]
    for row in tabela.rows:
        for i, cell in enumerate(row.cells):
            cell.width = larguras[i]


def _celula_quadrante(cell, quadrante, quadrante_ativo):
    nome = NOME_QUADRANTE[quadrante]
    ativo = (quadrante == quadrante_ativo)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome + ("  ◄ AQUI" if ativo else ""))
    run.font.size = Pt(9)
    run.font.bold = ativo
    if ativo:
        _sombrear_celula(cell, QUADRANTE_COR[quadrante])


def gerar_word_etapa8(estudo) -> BytesIO:
    """Gera o one-pager Word da Etapa 8. Retorna BytesIO pronto para download."""
    analise = estudo.estrategia_categoria or {}

    doc = Document()
    secao = doc.sections[0]
    secao.left_margin = Cm(1.8)
    secao.right_margin = Cm(1.8)
    secao.top_margin = Cm(1.5)
    secao.bottom_margin = Cm(1.5)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(10)

    titulo = doc.add_paragraph()
    run = titulo.add_run("Estratégia da Categoria — One-Pager")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM

    sub = doc.add_paragraph()
    quadrante = analise.get("quadrante", "")
    run_sub = sub.add_run(
        f"{analise.get('_categoria', 'Categoria')} · Matriz de Kraljic · "
        f"Quadrante: {NOME_QUADRANTE.get(quadrante, quadrante)}"
    )
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = COR_CINZA
    run_sub.italic = True

    _linha_divisoria(doc)

    # Matriz visual
    _heading(doc, "Posição na Matriz de Kraljic")
    _matriz_kraljic_word(doc, quadrante)
    eixos = doc.add_paragraph()
    run_eixos = eixos.add_run(
        f"Impacto financeiro: {analise.get('_impacto', '—')} ({analise.get('_origem_impacto', '—')})  ·  "
        f"Risco de suprimento: {analise.get('_risco', '—')} ({analise.get('_origem_risco', '—')})"
    )
    run_eixos.font.size = Pt(8)
    run_eixos.italic = True
    run_eixos.font.color.rgb = COR_CINZA

    if analise.get("resumo_posicao"):
        doc.add_paragraph(analise["resumo_posicao"])

    # Estratégia
    if analise.get("estrategia_recomendada"):
        _heading(doc, "Estratégia Recomendada")
        doc.add_paragraph(analise["estrategia_recomendada"])
        p = doc.add_paragraph()
        p.add_run(f"Fornecedores sugeridos: ").bold = True
        p.add_run(analise.get("numero_fornecedores_sugerido", "—"))
        p2 = doc.add_paragraph()
        p2.add_run(f"Relacionamento: ").bold = True
        p2.add_run(analise.get("tipo_relacionamento", "—"))

    # Ações táticas
    acoes = analise.get("acoes_taticas", [])
    if acoes:
        _heading(doc, "Ações Táticas")
        for a in acoes:
            doc.add_paragraph(
                f"[{a.get('prazo', '—')}] {a.get('acao', '')} — {a.get('racional', '')}",
                style="List Bullet",
            )

    # Árvore de categoria
    arvore = analise.get("arvore_categoria", {})
    if arvore:
        _heading(doc, "Árvore de Categoria")
        subs = ", ".join(arvore.get("subcategorias", [])) or "—"
        doc.add_paragraph(
            f"{arvore.get('macro', '—')} → {arvore.get('categoria', '—')} → {subs}"
        )

    # Alertas
    alertas = analise.get("alertas_estrategicos", [])
    if alertas:
        _heading(doc, "Alertas Estratégicos")
        for al in alertas:
            doc.add_paragraph(al, style="List Bullet")

    # Limitações
    faltantes = analise.get("faltantes", [])
    if faltantes:
        _heading(doc, "Limitações")
        for f in faltantes:
            doc.add_paragraph(f, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Geração Excel — mesmo estilo visual das Etapas 5 e 7
# ---------------------------------------------------------------------------

FILL_HEADER = PatternFill(start_color="1F3A5F", end_color="1F3A5F", fill_type="solid")
FONT_HEADER = Font(name="Arial", size=10, bold=True, color="FFFFFF")
FONT_BODY = Font(name="Arial", size=10)
FONT_BODY_BOLD = Font(name="Arial", size=10, bold=True)
THIN = Side(style="thin", color="CCCCCC")
BORDA = Border(top=THIN, bottom=THIN, left=THIN, right=THIN)


def _cabecalho_aba(ws, colunas):
    for i, texto in enumerate(colunas, start=1):
        cell = ws.cell(row=1, column=i, value=texto)
        cell.font = FONT_HEADER
        cell.fill = FILL_HEADER
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = BORDA
    ws.row_dimensions[1].height = 28
    ws.freeze_panes = "A2"


def gerar_excel_etapa8(estudo) -> BytesIO:
    """Gera o Excel da Etapa 8 (3 abas). Retorna BytesIO pronto para download."""
    analise = estudo.estrategia_categoria or {}

    wb = Workbook()
    _aba_posicionamento(wb, analise)
    _aba_acoes(wb, analise)
    _aba_arvore(wb, analise)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _aba_posicionamento(wb, analise):
    ws = wb.active
    ws.title = "Posicionamento Kraljic"
    _cabecalho_aba(ws, ["Dimensão", "Valor"])

    quadrante = analise.get("quadrante", "")
    linhas = [
        ("Categoria", analise.get("_categoria", "—")),
        ("Quadrante Kraljic", NOME_QUADRANTE.get(quadrante, quadrante)),
        ("Impacto financeiro", analise.get("_impacto", "—")),
        ("Origem do impacto", analise.get("_origem_impacto", "—")),
        ("Risco de suprimento", analise.get("_risco", "—")),
        ("Origem do risco", analise.get("_origem_risco", "—")),
        ("Posição (resumo)", analise.get("resumo_posicao", "—")),
        ("Estratégia recomendada", analise.get("estrategia_recomendada", "—")),
        ("Nº fornecedores sugerido", analise.get("numero_fornecedores_sugerido", "—")),
        ("Tipo de relacionamento", analise.get("tipo_relacionamento", "—")),
    ]
    for r, (dim, val) in enumerate(linhas, start=2):
        ws.cell(row=r, column=1, value=dim).font = FONT_BODY_BOLD
        ws.cell(row=r, column=2, value=val).font = FONT_BODY
        for c in (1, 2):
            cell = ws.cell(row=r, column=c)
            cell.border = BORDA
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 70


def _aba_acoes(wb, analise):
    ws = wb.create_sheet("Ações Táticas")
    _cabecalho_aba(ws, ["Prazo", "Ação", "Racional"])

    acoes = analise.get("acoes_taticas", [])
    for r, a in enumerate(acoes, start=2):
        ws.cell(row=r, column=1, value=a.get("prazo", "—")).font = FONT_BODY
        ws.cell(row=r, column=2, value=a.get("acao", "")).font = FONT_BODY
        ws.cell(row=r, column=3, value=a.get("racional", "")).font = FONT_BODY
        for c in (1, 2, 3):
            cell = ws.cell(row=r, column=c)
            cell.border = BORDA
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 45
    ws.column_dimensions["C"].width = 50


def _aba_arvore(wb, analise):
    ws = wb.create_sheet("Árvore e Alertas")
    arvore = analise.get("arvore_categoria", {})

    ws.cell(row=1, column=1, value="Árvore de Categoria").font = FONT_HEADER
    ws.cell(row=1, column=1).fill = FILL_HEADER
    ws.cell(row=1, column=2, value="").fill = FILL_HEADER
    ws.row_dimensions[1].height = 20

    linhas_arvore = [
        ("Macro-categoria", arvore.get("macro", "—")),
        ("Categoria", arvore.get("categoria", "—")),
        ("Subcategorias", ", ".join(arvore.get("subcategorias", [])) or "—"),
    ]
    r = 2
    for dim, val in linhas_arvore:
        ws.cell(row=r, column=1, value=dim).font = FONT_BODY_BOLD
        ws.cell(row=r, column=2, value=val).font = FONT_BODY
        for c in (1, 2):
            ws.cell(row=r, column=c).border = BORDA
            ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
        r += 1

    r += 1
    ws.cell(row=r, column=1, value="Alertas Estratégicos").font = FONT_HEADER
    ws.cell(row=r, column=1).fill = FILL_HEADER
    ws.cell(row=r, column=2, value="").fill = FILL_HEADER
    r += 1
    alertas = analise.get("alertas_estrategicos", [])
    if not alertas:
        ws.cell(row=r, column=1, value="—").font = FONT_BODY
        ws.cell(row=r, column=2, value="(nenhum registrado)").font = FONT_BODY
        r += 1
    else:
        for i, al in enumerate(alertas, start=1):
            ws.cell(row=r, column=1, value=f"#{i}").font = FONT_BODY_BOLD
            ws.cell(row=r, column=2, value=al).font = FONT_BODY
            for c in (1, 2):
                ws.cell(row=r, column=c).border = BORDA
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 70

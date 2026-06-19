"""
etapa7.py — Etapa 7: Recomendações Finais (NEUTRA).

O que faz:
1. É SÍNTESE FINAL — consolida Comparação Técnica (Etapa 5) × Equalização
   Comercial (Etapa 6) numa visão de decisão. NÃO reanalisa nada, só
   correlaciona o que já foi produzido.
2. NUNCA indica qual fornecedor escolher. Mostra savings no topo, os três
   "melhores" sob óticas diferentes (preço / técnica / custo-benefício),
   cenários de decisão com trade-offs explícitos, e pontos de negociação
   por fornecedor.
3. Com 1 fornecedor só: os três "melhores" colapsam no mesmo nome — sem
   fingir comparação. Mesmo tratamento dado à Etapa 5.
4. Grava em estudo.recomendacoes.
5. Gera Word (one-pager) e Excel (3 abas), seguindo o mesmo estilo visual
   da Etapa 5 para manter consistência entre os documentos do estudo.
   PPT na marca A&M: pendente layout do usuário (mesma regra da Etapa 5).
"""

import json
from io import BytesIO

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from ia import call_claude

MAX_TOKENS_ETAPA7 = 8000


# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------

SYSTEM_RECOMENDACOES = """
Você é um especialista sênior em procurement produzindo a SÍNTESE FINAL DE
DECISÃO de um estudo de propostas. Você NÃO está analisando documentos
brutos — está correlacionando duas análises que já foram feitas: a
comparação técnica (conformidade, gaps, escopo) e a equalização comercial
(preço equalizado, savings, on-tops). Sua função é dar ao consultor uma
visão de decisão SEM decidir por ele.

REGRA ABSOLUTA DE NEUTRALIDADE: você NUNCA recomenda qual fornecedor
escolher. Não use frases como "recomendamos X", "a melhor opção é X",
"sugerimos optar por X". Em vez disso, mostre os trade-offs de cada cenário
e deixe a decisão explícita para o consultor/cliente.

Você recebe: savings e preço equalizado por fornecedor (Etapa 6), e
conformidade/gaps/escopo por fornecedor (Etapa 5).

Responda SOMENTE com um objeto JSON válido, sem texto antes ou depois:

{
  "n_fornecedores": <número de fornecedores avaliados>,
  "eh_comparacao_real": <true se 2+ fornecedores, false se só 1>,
  "savings_destaque": {
    "maior_savings_absoluto": <número ou null>,
    "fornecedor_maior_savings": "<nome ou '—'>",
    "resumo": "<1-2 frases: onde está o maior savings e o range entre fornecedores>"
  },
  "tres_melhores": {
    "melhor_preco": {"fornecedor": "<nome>", "justificativa": "<1 frase, baseada em savings/preço equalizado>"},
    "melhor_tecnica": {"fornecedor": "<nome>", "justificativa": "<1 frase, baseada em conformidade/gaps>"},
    "melhor_custo_beneficio": {"fornecedor": "<nome>", "justificativa": "<1-2 frases explicando o raciocínio qualitativo entre preço e técnica — sem fórmula de peso numérico>"}
  },
  "cenarios_decisao": [
    {
      "nome": "<nome do cenário, ex.: 'Foco em economia', 'Foco em risco técnico zero', 'Meio-termo'>",
      "fornecedor_associado": "<nome>",
      "descricao": "<o que esse cenário prioriza>",
      "trade_off": "<o que se ganha e o que se perde escolhendo este cenário — explícito, sem julgar se vale a pena>"
    }
  ],
  "pontos_negociacao": [
    {
      "fornecedor": "<nome>",
      "alavancas": [
        {"ponto": "<alavanca concreta de negociação>", "origem": "gap_mandatorio|on_top_escopo|on_top_desvio|condicoes_comerciais", "argumento": "<1 frase: como usar isso na negociação>"}
      ]
    }
  ],
  "leitura_final": "<parágrafo: resumo equilibrado da situação, sem indicar caminho a seguir>",
  "premissas_registradas": ["<premissa assumida nesta síntese>"],
  "faltantes": ["<limitação desta síntese final>"]
}

Regras:
- Com 1 fornecedor só: eh_comparacao_real=false. tres_melhores ainda deve
  ser preenchido, mas os três campos terão o MESMO fornecedor (não invente
  um segundo). cenarios_decisao fica com 1 único cenário descrevendo a
  situação, sem comparação. Registre isso em faltantes.
- melhor_preco: baseie-se em savings_vs_baseline ou preco_total_equalizado
  da equalização comercial — não invente número novo.
- melhor_tecnica: baseie-se em quantidade/severidade de gaps mandatórios e
  conformidade — não invente score numérico.
- melhor_custo_beneficio: é leitura qualitativa, não fórmula. Explique o
  raciocínio (ex.: "segundo colocado em preço mas sem gaps mandatórios").
- pontos_negociacao.alavancas.origem deve refletir de onde veio o dado:
  gap_mandatorio (não cumpre mandatório), on_top_escopo (ajuste por delta
  do edital), on_top_desvio (inclusão/exclusão/desvio do fornecedor), ou
  condicoes_comerciais (frete/impostos/prazo/pagamento).
- NUNCA use linguagem de recomendação direta. Cenários e trade-offs, não
  conclusões.
- LIMITE DE TAMANHO: cenarios_decisao máximo 4 itens; pontos_negociacao
  máximo 1 por fornecedor com até 5 alavancas cada; premissas e faltantes
  máximo 6 itens cada. O JSON completo deve caber em 7000 tokens.
"""


# ---------------------------------------------------------------------------
# Montagem de contexto (lê do Estudo, não relê documentos brutos)
# ---------------------------------------------------------------------------

def _resumo_comparacao_tecnica_ctx(estudo) -> str:
    comp = estudo.comparacao_tecnica
    if not comp:
        return "Comparação técnica (Etapa 5): não disponível."
    linhas = ["COMPARAÇÃO TÉCNICA (Etapa 5):"]
    linhas.append(f"Resumo executivo: {comp.get('resumo_executivo','—')}")
    gaps = comp.get("gaps_mandatorios", [])
    if gaps:
        linhas.append("Gaps mandatórios por fornecedor:")
        for g in gaps:
            reqs = ", ".join(g.get("requisitos_nao_cumpridos", []))
            linhas.append(f"- {g.get('fornecedor','?')}: {reqs} — {g.get('leitura','')}")
    matriz = comp.get("matriz_requisitos", [])
    if matriz:
        n_cumpre, n_parcial, n_nao = {}, {}, {}
        for item in matriz:
            for s in item.get("status_por_fornecedor", []):
                forn = s.get("fornecedor", "?")
                status = s.get("status", "—")
                if status == "cumpre":
                    n_cumpre[forn] = n_cumpre.get(forn, 0) + 1
                elif status == "parcial":
                    n_parcial[forn] = n_parcial.get(forn, 0) + 1
                elif status == "não cumpre":
                    n_nao[forn] = n_nao.get(forn, 0) + 1
        fornecedores = set(n_cumpre) | set(n_parcial) | set(n_nao)
        for forn in fornecedores:
            linhas.append(
                f"- {forn}: {n_cumpre.get(forn,0)} cumpre, {n_parcial.get(forn,0)} parcial, "
                f"{n_nao.get(forn,0)} não cumpre (de {len(matriz)} requisitos)"
            )
    return "\n".join(linhas)


def _resumo_equalizacao_comercial_ctx(estudo) -> str:
    equal = estudo.equalizacao_comercial
    if not equal:
        return "Equalização comercial (Etapa 6): não disponível."
    linhas = ["EQUALIZAÇÃO COMERCIAL (Etapa 6):"]
    linhas.append(f"Moeda: {equal.get('moeda_referencia','—')} | Taxa de desconto aplicada: {equal.get('taxa_desconto_aplicada','—')}%")
    linhas.append(f"Síntese: {equal.get('sintese_comparativa','—')}")
    for forn in equal.get("por_fornecedor", []):
        linhas.append(
            f"- {forn.get('fornecedor','?')}: preço total equalizado = {forn.get('preco_total_equalizado')}, "
            f"savings vs baseline = {forn.get('savings_vs_baseline')} ({forn.get('savings_percentual')}%)"
        )
        on_tops_escopo = forn.get("on_tops_escopo", [])
        if on_tops_escopo:
            linhas.append(f"  on-tops de escopo: {json.dumps(on_tops_escopo, ensure_ascii=False)}")
        on_tops_desvio = forn.get("on_tops_desvio", [])
        if on_tops_desvio:
            linhas.append(f"  on-tops de desvio: {json.dumps(on_tops_desvio, ensure_ascii=False)}")
        ajustes = forn.get("ajustes_condicoes", [])
        if ajustes:
            linhas.append(f"  ajustes de condições: {json.dumps(ajustes, ensure_ascii=False)}")
    return "\n".join(linhas)


def _resumo_propostas_tecnicas_ctx(estudo) -> str:
    propostas = estudo.propostas_tecnicas or []
    if not propostas:
        return "Propostas técnicas (Etapa 4): nenhuma disponível."
    linhas = ["PROPOSTAS TÉCNICAS (Etapa 4) — flags de não-conformidade mandatória:"]
    for p in propostas:
        if p.get("nao_cumpre_mandatorio"):
            linhas.append(
                f"- {p.get('fornecedor','?')}: NÃO CUMPRE mandatório(s) "
                f"{', '.join(p.get('mandatorios_nao_cumpridos', []))}"
            )
    return "\n".join(linhas)


def _resumo_propostas_comerciais_ctx(estudo) -> str:
    propostas = estudo.propostas_comerciais or []
    if not propostas:
        return "Propostas comerciais (Etapa 4B): nenhuma disponível."
    linhas = ["CONDIÇÕES COMERCIAIS (Etapa 4B):"]
    for p in propostas:
        linhas.append(
            f"- {p.get('fornecedor','?')}: pagamento={p.get('condicoes_pagamento','—')}, "
            f"impostos_inclusos={p.get('impostos_inclusos','—')}, frete_incluso={p.get('frete_incluso','—')}"
        )
    return "\n".join(linhas)


def _parse_resposta(resposta_bruta: str) -> dict:
    """Parser com fallback de truncamento (mesmo padrão das outras etapas)."""
    texto = resposta_bruta.strip()
    if texto.startswith("```"):
        linhas = texto.split("\n")
        texto = "\n".join(linhas[1:-1]).strip()
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        texto_cortado = texto.rstrip().rstrip(",")
        abertas = texto_cortado.count("{") - texto_cortado.count("}")
        colchetes = texto_cortado.count("[") - texto_cortado.count("]")
        fechamento = "]" * colchetes + "}" * abertas
        try:
            return json.loads(texto_cortado + fechamento)
        except json.JSONDecodeError:
            raise ValueError(
                "JSON inválido mesmo após tentativa de correção. "
                "Estudo com muitos fornecedores — verifique o volume de dados."
            )


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def rodar_etapa7(estudo) -> dict:
    """
    Executa a Etapa 7 completa.

    Retorna dict com:
        - tem_dados : bool
        - analise   : dict completo (JSON do Claude) ou None
        - resumo    : texto legível pra UI
    """
    if not estudo.comparacao_tecnica and not estudo.equalizacao_comercial:
        msg = "Nem comparação técnica (Etapa 5) nem equalização comercial (Etapa 6) disponíveis — Etapa 7 não pôde rodar."
        estudo.add_faltante(msg)
        estudo.recomendacoes = None
        return {"tem_dados": False, "analise": None, "resumo": f"⚠️ {msg}"}

    contexto = (
        f"Categoria: {estudo.categoria}\n"
        f"Micro-categoria: {estudo.micro_categoria or '—'}\n\n"
        f"{_resumo_comparacao_tecnica_ctx(estudo)}\n\n"
        f"{_resumo_equalizacao_comercial_ctx(estudo)}\n\n"
        f"{_resumo_propostas_tecnicas_ctx(estudo)}\n\n"
        f"{_resumo_propostas_comerciais_ctx(estudo)}"
    )

    with st.spinner("Etapa 7 — montando recomendações finais (neutras)..."):
        resposta_bruta = call_claude(
            messages=[{"role": "user", "content": contexto}],
            system=SYSTEM_RECOMENDACOES,
            max_tokens=MAX_TOKENS_ETAPA7,
        )

    try:
        analise = _parse_resposta(resposta_bruta)
    except (json.JSONDecodeError, ValueError) as e:
        st.error(f"Erro ao interpretar resposta da IA na Etapa 7: {e}\n\nResposta bruta:\n{resposta_bruta}")
        st.stop()

    estudo.recomendacoes = analise

    for p in analise.get("premissas_registradas", []):
        estudo.add_premissa(p)
    for f in analise.get("faltantes", []):
        estudo.add_faltante(f)
    if not analise.get("eh_comparacao_real", True):
        estudo.add_faltante(
            "Etapa 7 rodou com apenas 1 fornecedor — os 'três melhores' colapsam "
            "no mesmo nome, sem comparação real de cenários."
        )

    estudo.etapa_atual = 7

    resumo = _montar_resumo(analise)
    return {"tem_dados": True, "analise": analise, "resumo": resumo}


def _montar_resumo(analise: dict) -> str:
    linhas = []

    savings = analise.get("savings_destaque", {})
    if savings:
        linhas.append(
            f"**💰 Savings em destaque:** {savings.get('resumo','—')}"
        )

    if not analise.get("eh_comparacao_real", True):
        linhas.append("\n⚠️ _Apenas 1 fornecedor — sem comparação real entre cenários._")

    tres = analise.get("tres_melhores", {})
    if tres:
        linhas.append("\n**Os três 'melhores' (óticas diferentes):**")
        mp = tres.get("melhor_preco", {})
        mt = tres.get("melhor_tecnica", {})
        mcb = tres.get("melhor_custo_beneficio", {})
        linhas.append(f"- **Melhor preço:** {mp.get('fornecedor','—')} — {mp.get('justificativa','')}")
        linhas.append(f"- **Melhor técnica:** {mt.get('fornecedor','—')} — {mt.get('justificativa','')}")
        linhas.append(f"- **Melhor custo-benefício:** {mcb.get('fornecedor','—')} — {mcb.get('justificativa','')}")

    cenarios = analise.get("cenarios_decisao", [])
    if cenarios:
        linhas.append("\n**Cenários de decisão:**")
        for c in cenarios:
            linhas.append(f"\n_{c.get('nome','?')}_ ({c.get('fornecedor_associado','—')})")
            linhas.append(f"{c.get('descricao','')}")
            linhas.append(f"Trade-off: {c.get('trade_off','')}")

    pontos = analise.get("pontos_negociacao", [])
    if pontos:
        linhas.append("\n**Pontos de negociação:**")
        for p in pontos:
            linhas.append(f"\n**{p.get('fornecedor','?')}:**")
            for a in p.get("alavancas", []):
                linhas.append(f"- [{a.get('origem','?')}] {a.get('ponto','')} — {a.get('argumento','')}")

    leitura = analise.get("leitura_final")
    if leitura:
        linhas.append(f"\n**Leitura final:**\n{leitura}")

    faltantes = analise.get("faltantes", [])
    if faltantes:
        linhas.append("\n**Limitações desta síntese:**")
        for f in faltantes:
            linhas.append(f"- {f}")

    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Geração Word (one-pager) — mesmo estilo visual da Etapa 5
# ---------------------------------------------------------------------------

COR_AZUL_AM = RGBColor(0x1F, 0x3A, 0x5F)
COR_CINZA = RGBColor(0x59, 0x59, 0x59)
COR_VERDE_SAVINGS = RGBColor(0x1E, 0x7A, 0x3C)


def _heading(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM


def _linha_divisoria(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _sombrear_celula(cell, cor_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


def gerar_word_etapa7(estudo) -> BytesIO:
    """Gera o one-pager Word da Etapa 7. Retorna BytesIO pronto para download."""
    analise = estudo.recomendacoes or {}

    doc = Document()

    secao = doc.sections[0]
    secao.left_margin = Cm(1.8)
    secao.right_margin = Cm(1.8)
    secao.top_margin = Cm(1.5)
    secao.bottom_margin = Cm(1.5)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(10)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = titulo.add_run("Recomendações Finais — One-Pager")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM

    sub = doc.add_paragraph()
    run_sub = sub.add_run(
        f"{estudo.micro_categoria or estudo.categoria or 'Categoria não identificada'} · "
        f"{analise.get('n_fornecedores', '—')} fornecedor(es) · síntese neutra, sem indicação de escolha"
    )
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = COR_CINZA
    run_sub.italic = True

    _linha_divisoria(doc)

    if not analise.get("eh_comparacao_real", True):
        aviso = doc.add_paragraph()
        run_aviso = aviso.add_run(
            "⚠ Apenas 1 fornecedor avaliado — os cenários abaixo não representam "
            "comparação entre fornecedores."
        )
        run_aviso.font.size = Pt(9)
        run_aviso.italic = True
        run_aviso.font.color.rgb = COR_CINZA

    # Savings em destaque
    savings = analise.get("savings_destaque", {})
    if savings:
        _heading(doc, "Savings em Destaque")
        p = doc.add_paragraph()
        valor = savings.get("maior_savings_absoluto")
        if valor is not None:
            run_v = p.add_run(f"{valor:,.2f} ")
            run_v.font.size = Pt(16)
            run_v.font.bold = True
            run_v.font.color.rgb = COR_VERDE_SAVINGS
            p.add_run(f"({savings.get('fornecedor_maior_savings','—')})").font.size = Pt(10)
        doc.add_paragraph(savings.get("resumo", ""))

    # Três melhores
    tres = analise.get("tres_melhores", {})
    if tres:
        _heading(doc, "Os Três 'Melhores' (Óticas Diferentes)")
        tabela = doc.add_table(rows=1, cols=3)
        tabela.style = "Table Grid"
        tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tabela.rows[0].cells
        for i, titulo_col in enumerate(["Melhor Preço", "Melhor Técnica", "Melhor Custo-Benefício"]):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(titulo_col)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _sombrear_celula(hdr[i], "1F3A5F")

        row = tabela.add_row().cells
        chaves = ["melhor_preco", "melhor_tecnica", "melhor_custo_beneficio"]
        for i, chave in enumerate(chaves):
            item = tres.get(chave, {})
            row[i].text = ""
            p_nome = row[i].paragraphs[0]
            run_nome = p_nome.add_run(item.get("fornecedor", "—"))
            run_nome.font.bold = True
            run_nome.font.size = Pt(9)
            p_just = row[i].add_paragraph()
            run_just = p_just.add_run(item.get("justificativa", ""))
            run_just.font.size = Pt(8)

        larguras = [Cm(5.5)] * 3
        for row in tabela.rows:
            for i, cell in enumerate(row.cells):
                cell.width = larguras[i]

    # Cenários de decisão
    cenarios = analise.get("cenarios_decisao", [])
    if cenarios:
        _heading(doc, "Cenários de Decisão")
        for c in cenarios:
            p = doc.add_paragraph()
            run_nome = p.add_run(f"{c.get('nome','?')} ")
            run_nome.bold = True
            run_nome.font.size = Pt(10)
            run_forn = p.add_run(f"({c.get('fornecedor_associado','—')})")
            run_forn.italic = True
            run_forn.font.size = Pt(9)
            run_forn.font.color.rgb = COR_CINZA
            doc.add_paragraph(c.get("descricao", ""), style="List Bullet")
            p_trade = doc.add_paragraph()
            run_label = p_trade.add_run("Trade-off: ")
            run_label.bold = True
            run_label.font.size = Pt(9)
            p_trade.add_run(c.get("trade_off", "")).font.size = Pt(9)

    # Pontos de negociação
    pontos = analise.get("pontos_negociacao", [])
    if pontos:
        _heading(doc, "Pontos de Negociação")
        for p_item in pontos:
            p = doc.add_paragraph()
            run_forn = p.add_run(p_item.get("fornecedor", "?"))
            run_forn.bold = True
            run_forn.font.size = Pt(10)
            for a in p_item.get("alavancas", []):
                doc.add_paragraph(
                    f"[{a.get('origem','?')}] {a.get('ponto','')} — {a.get('argumento','')}",
                    style="List Bullet",
                )

    # Leitura final
    if analise.get("leitura_final"):
        _heading(doc, "Leitura Final")
        doc.add_paragraph(analise["leitura_final"])

    # Limitações
    faltantes = analise.get("faltantes", [])
    if faltantes:
        _heading(doc, "Limitações desta Síntese")
        for f in faltantes:
            doc.add_paragraph(f, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Geração Excel (3 abas) — mesmo estilo visual da Etapa 5
# ---------------------------------------------------------------------------

FILL_HEADER = PatternFill(start_color="1F3A5F", end_color="1F3A5F", fill_type="solid")
FONT_HEADER = Font(name="Arial", size=10, bold=True, color="FFFFFF")
FONT_BODY = Font(name="Arial", size=10)
FONT_BODY_BOLD = Font(name="Arial", size=10, bold=True)
THIN = Side(style="thin", color="CCCCCC")
BORDA = Border(top=THIN, bottom=THIN, left=THIN, right=THIN)
FILL_DESTAQUE = PatternFill(start_color="DCE6F0", end_color="DCE6F0", fill_type="solid")


def _cabecalho_aba(ws, colunas):
    for i, texto in enumerate(colunas, start=1):
        cell = ws.cell(row=1, column=i, value=texto)
        cell.font = FONT_HEADER
        cell.fill = FILL_HEADER
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = BORDA
    ws.row_dimensions[1].height = 28
    ws.freeze_panes = "A2"


def gerar_excel_etapa7(estudo) -> BytesIO:
    """Gera o Excel da Etapa 7 (3 abas). Retorna BytesIO pronto para download."""
    analise = estudo.recomendacoes or {}
    equal = estudo.equalizacao_comercial or {}

    wb = Workbook()
    _aba_resumo_comparativo(wb, analise, equal)
    _aba_cenarios(wb, analise)
    _aba_negociacao(wb, analise)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _aba_resumo_comparativo(wb, analise, equal):
    ws = wb.active
    ws.title = "Resumo Comparativo"
    _cabecalho_aba(ws, ["Fornecedor", "Preço Equalizado", "Savings vs Baseline", "Savings %", "Destaque"])

    tres = analise.get("tres_melhores", {})
    destaques_por_fornecedor = {}
    for chave, label in [
        ("melhor_preco", "Melhor preço"),
        ("melhor_tecnica", "Melhor técnica"),
        ("melhor_custo_beneficio", "Melhor custo-benefício"),
    ]:
        forn = tres.get(chave, {}).get("fornecedor")
        if forn:
            destaques_por_fornecedor.setdefault(forn, []).append(label)

    por_fornecedor = equal.get("por_fornecedor", [])
    for r, forn in enumerate(por_fornecedor, start=2):
        nome = forn.get("fornecedor", "?")
        ws.cell(row=r, column=1, value=nome).font = FONT_BODY_BOLD
        ws.cell(row=r, column=2, value=forn.get("preco_total_equalizado")).font = FONT_BODY
        ws.cell(row=r, column=3, value=forn.get("savings_vs_baseline")).font = FONT_BODY
        ws.cell(row=r, column=4, value=forn.get("savings_percentual")).font = FONT_BODY
        destaque_texto = ", ".join(destaques_por_fornecedor.get(nome, []))
        cell_destaque = ws.cell(row=r, column=5, value=destaque_texto)
        cell_destaque.font = FONT_BODY
        if destaque_texto:
            cell_destaque.fill = FILL_DESTAQUE
        for c in range(1, 6):
            ws.cell(row=r, column=c).border = BORDA

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 35


def _aba_cenarios(wb, analise):
    ws = wb.create_sheet("Cenários de Decisão")
    _cabecalho_aba(ws, ["Cenário", "Fornecedor Associado", "Descrição", "Trade-off"])

    cenarios = analise.get("cenarios_decisao", [])
    for r, c in enumerate(cenarios, start=2):
        ws.cell(row=r, column=1, value=c.get("nome", "?")).font = FONT_BODY_BOLD
        ws.cell(row=r, column=2, value=c.get("fornecedor_associado", "—")).font = FONT_BODY
        ws.cell(row=r, column=3, value=c.get("descricao", "")).font = FONT_BODY
        ws.cell(row=r, column=4, value=c.get("trade_off", "")).font = FONT_BODY
        for col in range(1, 5):
            cell = ws.cell(row=r, column=col)
            cell.border = BORDA
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 45
    ws.column_dimensions["D"].width = 45


def _aba_negociacao(wb, analise):
    ws = wb.create_sheet("Pontos de Negociação")
    _cabecalho_aba(ws, ["Fornecedor", "Origem", "Ponto", "Argumento"])

    pontos = analise.get("pontos_negociacao", [])
    r = 2
    for p in pontos:
        nome = p.get("fornecedor", "?")
        alavancas = p.get("alavancas", [])
        if not alavancas:
            ws.cell(row=r, column=1, value=nome).font = FONT_BODY_BOLD
            ws.cell(row=r, column=2, value="—").font = FONT_BODY
            ws.cell(row=r, column=3, value="(nenhum ponto registrado)").font = FONT_BODY
            ws.cell(row=r, column=4, value="").font = FONT_BODY
            for col in range(1, 5):
                ws.cell(row=r, column=col).border = BORDA
            r += 1
            continue
        for a in alavancas:
            ws.cell(row=r, column=1, value=nome).font = FONT_BODY_BOLD
            ws.cell(row=r, column=2, value=a.get("origem", "—")).font = FONT_BODY
            ws.cell(row=r, column=3, value=a.get("ponto", "")).font = FONT_BODY
            ws.cell(row=r, column=4, value=a.get("argumento", "")).font = FONT_BODY
            for col in range(1, 5):
                cell = ws.cell(row=r, column=col)
                cell.border = BORDA
                cell.alignment = Alignment(wrap_text=True, vertical="top")
            r += 1

    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 20
    ws.column_dimensions["C"].width = 35
    ws.column_dimensions["D"].width = 45

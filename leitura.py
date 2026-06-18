"""
leitura.py — Leitura de arquivos.

Recebe um arquivo enviado pelo usuário (PDF, Word, Excel ou texto) e
devolve o conteúdo dele como texto, pronto pra ser analisado pela IA.

Portado do app.py atual, sem mudança de comportamento — é código já
testado em produção, só reorganizado num módulo próprio.
"""

import pandas as pd
import pdfplumber
from docx import Document


def read_pdf(file) -> str:
    """Extrai texto de PDF (inclui tabelas, comuns em propostas)."""
    try:
        text_parts = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                # Tabelas em PDFs comuns em propostas
                for table in page.extract_tables():
                    for row in table:
                        if row:
                            text_parts.append(" | ".join([str(c) if c else "" for c in row]))
        return "\n".join(text_parts)
    except Exception as e:
        return f"[ERRO ao ler PDF: {e}]"


def read_docx(file) -> str:
    """Extrai texto de DOCX (Word), incluindo tabelas."""
    try:
        doc = Document(file)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[ERRO ao ler DOCX: {e}]"


def read_xlsx(file) -> str:
    """Extrai conteúdo de Excel como texto estruturado (aba por aba)."""
    try:
        xl = pd.ExcelFile(file)
        parts = []
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(file, sheet_name=sheet_name, header=None)
            df = df.dropna(how="all").dropna(axis=1, how="all")
            if df.empty:
                continue
            parts.append(f"\n=== Aba: {sheet_name} ===")
            for _, row in df.iterrows():
                row_values = [str(v) for v in row.values if pd.notna(v) and str(v).strip()]
                if row_values:
                    parts.append(" | ".join(row_values))
        return "\n".join(parts)
    except Exception as e:
        return f"[ERRO ao ler Excel: {e}]"


def read_file(file) -> str:
    """Detecta o tipo do arquivo pelo nome e extrai o texto."""
    filename = file.name.lower()
    if filename.endswith(".pdf"):
        return read_pdf(file)
    elif filename.endswith(".docx"):
        return read_docx(file)
    elif filename.endswith((".xlsx", ".xls", ".xlsm")):
        return read_xlsx(file)
    elif filename.endswith(".txt"):
        return file.read().decode("utf-8", errors="ignore")
    else:
        return f"[Formato não suportado: {filename}]"
